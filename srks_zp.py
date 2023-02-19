# https://dashboard.render.com/web/srv-cfdsq7pgp3jolcmp3bg0
import os, json, requests, re, time, pandas
from urllib.parse import urlencode
import datetime
import docx
from io import BytesIO
from flask import Flask, request
 
TG_TOKEN = os.environ.get('TG_TOKEN')
B24_CLIENT, B24_SECRET = os.environ.get('B24_CLIENT'), os.environ.get('B24_SECRET')
ACT_LINK, CONTRACT_LINK = os.environ.get('ACT_LINK'), os.environ.get('CONTRACT_LINK')
RENDER_LINK = os.environ.get('RENDER_LINK')
app = Flask(__name__)


@app.route("/", methods=["GET", "POST"])
def get_message():    
    try:
        # if this isnt message
        if not 'message' in request.json:
            return "ok", 200
        
        user_list = b24.get_user_by_tg(request.json["message"]["from"]["username"])
        # if no 1 users with username in b24
        if len(user_list) == 0:
            requests.post(f"https://api.telegram.org/bot{TG_TOKEN}/{'sendMessage'}", data={'chat_id': request.json["message"]["chat"]["id"], 'parse_mode':'MarkdownV2', 
                          'text': 'в Б24 не найдено пользователя с вашим логином телеграм\. Добавьте его в [ваш профиль](https://zapovednik.bitrix24.ru/company/personal/user/).'})
            return "ok", 200 
        # if 2+ users
        elif len(user_list) > 1:
            requests.post(f"https://api.telegram.org/bot{TG_TOKEN}/{'sendMessage'}", data={'chat_id': request.json["message"]["chat"]["id"], 'parse_mode':'MarkdownV2', 
                          'text': f'в Б24 найдено несколько пользователей с вашим логином телеграм\. Получение задач невозможно, пока не останется лишь один\.' +\
                          str([f'[{user["ID"]}](https://zapovednik.bitrix24.ru/company/personal/user/{user["ID"]}/)' for user in user_list])})
            return "ok", 200
          
        # if start
        if request.json["message"]["text"]=="/start" or request.json["message"]["text"]=="/restart":
            requests.post(f"https://api.telegram.org/bot{TG_TOKEN}/sendMessage", data={'chat_id': request.json["message"]["chat"]["id"], 'parse_mode':'MarkdownV2', 
                          'text': '👋 я \- бот заполняющий шаблоны [Акта]({}) и [Контракта]({}) на основании переданных параметров\. Данные нигде не хранятся, кроме этого чата \(и оперативной памяти серверов render\.com в момент заполнения\)\. Для заполнения шаблона отправьте параметры в виде\: \n\n \{{"\_Итог\_": 50000\.00, "\_Адрес\_": "адрес регистрации", "\_Паспорт\_": "паспортные данные", "\_Банк\_": "АО «Тинькоф\-банк»", "\_ИНН\_": "инн", "\_СНИЛС\_": "снилс", "\_Счет\_": "банк\.счет", "\_КС\_": "корр\-счет", "\_БИК\_": "БИК банка"\}}'.format(ACT_LINK, CONTRACT_LINK)}) 
            return "ok", 200
          
        params = json.loads(request.json["message"]["text"])
        user = user_list[0]
        tasks = b24.get_tasks(user["ID"], params)
        params["tasks"] = tasks
        
        today      = datetime.datetime.today()
        last_month = (today - datetime.timedelta(days=29)).replace(day=1)
        month_dict = {'1': 'января', '2': 'февраля', '3': 'марта', '4': 'апреля', '5': 'мая', '6': 'июня', '7': 'июля', '8': 'августа', '9': 'сентябся', 10: 'октября', '11': 'ноября','12': 'декабря'}
    
        # Параметры по-умолчанию
        params = params | {'_ТЗДатаФормат_': f'«01» {month_dict[last_month.month.__str__()]} {last_month.year.__str__()}', '_ТЗДата_': last_month.strftime("%d.%m.%Y"), '_ТЗКонецМесяцаДата_': (today - datetime.timedelta(days=today.day)).strftime("%d.%m.%Y"),
                  '_ФИО_': '{} {} {}'.format(user['LAST_NAME'], user['NAME'], user['SECOND_NAME']), '_ФамилияИО_': '{} {}.{}.'.format(user['LAST_NAME'], user['NAME'][0], user['SECOND_NAME'][0] if len(user['SECOND_NAME']) else '')}
        params["_Итог_"] = str(params["_Итог_"]) + " ("+requests.get(f'https://htmlweb.ru/api/convert/num2str?num={params["_Итог_"]}&noLimit&html&uc=1').text+")"
        for k,v in {'act': 'https://disk.yandex.ru/i/49uSpSlpp1KGGA', 'contract': 'https://disk.yandex.ru/i/sLtvDNmMycOpkA'}.items():
            doc_link = b24.get_yandex_link(v)
            doc = docx.Document(BytesIO(requests.get(doc_link).content))
            doc = fill_doc(doc, params)
            doc_buf = BytesIO() 
            doc.save(doc_buf)
            _=requests.post(f"https://api.telegram.org/bot{TG_TOKEN}/sendDocument",
                       params={'chat_id': request.json["message"]["chat"]["id"]}, files={'document': (f'{k}.docx', doc_buf.getvalue())}
                       )
            _.status_code, _.text
        return "ok", 200
    except Exception as e:
        requests.post(f"https://api.telegram.org/bot{TG_TOKEN}/{'sendMessage'}", data={'chat_id': request.json["message"]["chat"]["id"], 'text': f'Возникла ошибка. Скиньте обезличенные данные, разберусь\n\n{e}'})
        return "ok", 200

class b24:
    def __init__(self):
        # import time, requests, json, pandas
        self.B24_CLIENT, self.B24_SECRET = B24_CLIENT, B24_SECRET 
        response = requests.get(f'https://{os.environ.get("B24_LOGIN")}:{os.environ.get("B24_PASSWORD")}@zapovednik.bitrix24.ru/oauth/authorize/?client_id={self.B24_CLIENT}')
        B24_CODE = re.search(f'http.+code=(.+?)[&]', response.url).groups(1)[0]
        response = requests.get(f'https://oauth.bitrix.info/oauth/token/?grant_type=authorization_code&client_id={self.B24_CLIENT}&client_secret={self.B24_SECRET}&code={B24_CODE}')
        self.B24_CRED = json.loads(response.text)
    # https://dev.1c-bitrix.ru/rest_help/tasks/task/tasks/tasks_task_list.php
    def call(self, method, params={}):
        # authorize
        if self.B24_CRED == None or self.B24_CRED['expires'] < time.time() + 10:
            self.__init__()
        response = requests.get(f'https://zapovednik.bitrix24.ru/rest/{method}?auth={self.B24_CRED["access_token"]}', params=params)
        return json.loads(response.text)["result"]
    def get_tasks(self, id_b24, params, CLOSED_DATE=None, tasks_head = 8):
        tasks = pandas.DataFrame.from_dict(self.call('tasks.task.list', params={'select[0]': 'TITLE', 'select[1]': 'DATE_START', 'select[2]': 'CLOSED_DATE', 'filter[REAL_STATUS]': 5, 'filter[<=CLOSED_DATE]': time.strftime("%d.%m.%Y", CLOSED_DATE if CLOSED_DATE!=None else time.gmtime()), 'filter[RESPONSIBLE_ID]': id_b24, 'order[DATE_START]': 'desc'})["tasks"])
        for column in tasks.columns: 
            if column.lower().find('date') > -1:
                tasks[column] = pandas.to_datetime(tasks[column]) # str2date
        tasks["hours"] = (tasks.closedDate-tasks.dateStart).astype('timedelta64[h]') # evaluate hours
        tasks.sort_values("hours", ascending=False, inplace=True) # sort by hours
        tasks = tasks.head(tasks_head).reset_index() # select N first rows
        tasks["total"] = (params["_Итог_"]*tasks.hours / tasks.hours.sum()).round(2) # evaluate rub from hours
        tasks.at[tasks_head-1, "total"] -= (tasks.total.sum() - params["_Итог_"]).round(2) # resolve cent problem in last row
        return tasks
    def get_user_by_tg(self, username: str) -> list:
        return self.call('user.get', {'filter[UF_USR_1676315711359]': username})
    # https://ru.stackoverflow.com/questions/1088300/как-скачивать-файлы-с-яндекс-диска
    def get_yandex_link(self, public_key):
        base_url  = 'https://cloud-api.yandex.net/v1/disk/public/resources/download?'
        final_url = base_url + urlencode(dict(public_key=public_key))
        response  = requests.get(final_url)
        return response.json()['href']

def fill_doc(doc, params):
    # rFonts = doc.styles['Normal'].element.rPr.rFonts
    # rFonts.set(docx.oxml.ns.qn("w:asciiTheme"), "Times New Roman")
    if type(docx) == str:
        doc = docx.Document(docx)
    font = doc.styles['Normal'].font
    font.name = 'Arial'
    # font.size = docx.shared.Pt(12)

    # replace table text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for k,v in params.items(): 
                            if k in run.text:
                                run.text = run.text.replace(k, str(v))

    # replace text
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for k,v in params.items():
                if k in run.text:
                    run.text = run.text.replace(k, str(v))
     
    # fill table https://python-docx.readthedocs.io/en/latest/api/table.html?highlight=add_row#docx.table._Cell
    tab = doc.tables[2]
    for index,task in params["tasks"].iterrows():
        _row = tab.add_row()
        _row.cells[0].text = str(index+1)
        _row.cells[1].text = task.title
        _row.cells[2].text = str(task.total)
    
    return doc

requests.get(f"https://api.telegram.org/bot{TG_TOKEN}/deleteWebhook") 
requests.post(f"https://api.telegram.org/bot{TG_TOKEN}/setWebhook", data={'url': RENDER_LINK})  
b24 = b24()
